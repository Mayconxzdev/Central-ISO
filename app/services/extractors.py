from __future__ import annotations

from pathlib import Path

from .office_protected import readable_office_path
from .paths import filesystem_path


def _safe_text(value: object) -> str:
    if value is None:
        return ""
    return str(value).strip()


def extract_text(path: Path) -> tuple[str, str]:
    """Return (text, status). Never executes macros or embedded code."""
    ext = path.suffix.lower()
    try:
        if ext in {".txt", ".csv", ".md", ".log"}:
            for encoding in ("utf-8", "cp1252", "latin-1"):
                try:
                    with open(filesystem_path(path), encoding=encoding, errors="strict") as stream:
                        return stream.read(), "lido"
                except UnicodeDecodeError:
                    continue
            with open(filesystem_path(path), encoding="utf-8", errors="replace") as stream:
                return stream.read(), "lido parcialmente"

        if ext == ".pdf":
            import fitz

            doc = fitz.open(filesystem_path(path))
            pages = []
            for idx, page in enumerate(doc):
                text = page.get_text("text").strip()
                if text:
                    pages.append(f"[Página {idx + 1}]\n{text}")
            doc.close()
            if pages:
                return "\n\n".join(pages), "lido"
            return "", "sem texto — OCR necessário"

        if ext == ".docx":
            from docx import Document

            with readable_office_path(path) as (readable, protection_status):
                if protection_status.startswith("protegido, senha") or protection_status.startswith("falha"):
                    return "", protection_status
                doc = Document(filesystem_path(readable))
                parts: list[str] = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        parts.append(paragraph.text.strip())
                for table_index, table in enumerate(doc.tables, start=1):
                    parts.append(f"[Tabela {table_index}]")
                    for row in table.rows:
                        parts.append(" | ".join(_safe_text(cell.text) for cell in row.cells))
            return "\n".join(parts), "lido" if protection_status == "nao protegido" else protection_status

        if ext in {".xlsx", ".xlsm"}:
            from openpyxl import load_workbook

            with readable_office_path(path) as (readable, protection_status):
                if protection_status.startswith("protegido, senha") or protection_status.startswith("falha"):
                    return "", protection_status
                wb = load_workbook(filesystem_path(readable), read_only=True, data_only=True, keep_vba=False)
                parts: list[str] = []
                for ws in wb.worksheets:
                    parts.append(f"[Aba: {ws.title}]")
                    for row in ws.iter_rows(values_only=True):
                        values = [_safe_text(value) for value in row]
                        if any(values):
                            parts.append(" | ".join(values))
                wb.close()
            return "\n".join(parts), "lido" if protection_status == "nao protegido" else protection_status

        if ext == ".odt":
            from odf import text as odf_text
            from odf.opendocument import load

            doc = load(filesystem_path(path))
            parts = []
            for node in doc.getElementsByType(odf_text.P):
                value = "".join(child.data for child in node.childNodes if hasattr(child, "data"))
                if value.strip():
                    parts.append(value.strip())
            return "\n".join(parts), "lido"

        return "", "formato não suportado"
    except PermissionError:
        return "", "protegido"
    except Exception as exc:  # noqa: BLE001
        message = str(exc).lower()
        if "password" in message or "encrypted" in message or "cript" in message:
            return "", "protegido"
        return f"ERRO: {exc}", "falha"
